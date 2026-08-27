from pathlib import Path
import math
import zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:
    Image = ImageDraw = ImageFont = None


OUT_DIR = Path(__file__).resolve().parent
PDF_OUT_DIR = OUT_DIR.parent / "output" / "pdf"
DOCX_PATH = OUT_DIR / "typing_certificate_intermediate.docx"
FINAL_PATH = OUT_DIR / "typing_certificate_intermediate_award.docx"
PDF_PATH = PDF_OUT_DIR / "typing_certificate_intermediate_award.pdf"
SEAL_PATH = OUT_DIR / "typing_certificate_seal.png"

FONT_SERIF = "Yu Mincho"
FONT_SANS = "Yu Gothic"
GOLD = "B99733"
DARK = "2B2116"
RED = "A2352A"
PALE_GOLD = "FBF6E5"


def set_font(run, name=FONT_SERIF, size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell, size=10, color=GOLD):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tc_pr.append(borders)


def paragraph_border(paragraph, position="bottom", color=GOLD, size=10, space=5):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{position}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def create_seal():
    if Image is None:
        return None
    size = 520
    image = Image.new("RGBA", (size, size), (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)
    center = size / 2
    outer = []
    points = 64
    for i in range(points * 2):
        angle = math.pi * i / points
        radius = 236 if i % 2 == 0 else 206
        outer.append((center + math.cos(angle) * radius, center + math.sin(angle) * radius))
    draw.polygon(outer, fill=(185, 151, 51, 255), outline=(121, 91, 18, 255))
    draw.ellipse((74, 74, 446, 446), fill=(242, 210, 104, 255), outline=(121, 91, 18, 255), width=8)
    draw.ellipse((118, 118, 402, 402), outline=(255, 250, 220, 255), width=6)
    try:
        font_big = ImageFont.truetype("C:/Windows/Fonts/yumin.ttf", 74)
        font_small = ImageFont.truetype("C:/Windows/Fonts/yugothb.ttc", 34)
    except Exception:
        font_big = font_small = None
    for text, y, font in [("中級", 194, font_big), ("合格", 280, font_big), ("TYPING", 134, font_small)]:
        bbox = draw.textbbox((0, 0), text, font=font)
        draw.text(((size - (bbox[2] - bbox[0])) / 2, y), text, fill=(82, 47, 8, 255), font=font)
    image.save(SEAL_PATH)
    return SEAL_PATH


def add_centered(paragraph, text, size, color=DARK, bold=False, font=FONT_SERIF, after=0, before=0):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=before, after=after, line=1.0)
    run = paragraph.add_run(text)
    set_font(run, font, size=size, color=color, bold=bold)
    return run


def build_docx():
    seal = create_seal()
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)

    style = doc.styles["Normal"]
    style.font.name = FONT_SERIF
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SERIF)
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor.from_string(DARK)

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(26.7)
    cell = table.cell(0, 0)
    cell.width = Cm(26.7)
    set_cell_borders(cell, size=16, color=GOLD)
    set_cell_shading(cell, PALE_GOLD)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, value in [("top", "560"), ("bottom", "500"), ("start", "780"), ("end", "780")]:
        margin = OxmlElement(f"w:{side}")
        margin.set(qn("w:w"), value)
        margin.set(qn("w:type"), "dxa")
        tc_mar.append(margin)
    tc_pr.append(tc_mar)

    p = cell.paragraphs[0]
    add_centered(p, "賞　状", 42, color=DARK, bold=True, after=6)

    p = cell.add_paragraph()
    add_centered(p, "タイピング検定　中級", 20, color=GOLD, bold=True, font=FONT_SANS, after=18)

    p = cell.add_paragraph()
    add_centered(p, "受賞者名", 13, color="6E6255", font=FONT_SANS, after=0)

    p = cell.add_paragraph()
    add_centered(p, "　　　　　　　　　　　　　　殿", 22, color=DARK, bold=True, after=2)
    paragraph_border(p, "bottom", color=GOLD, size=8, space=4)

    p = cell.add_paragraph()
    add_centered(
        p,
        "あなたはタイピング検定 中級において、正確で安定した入力技術と継続した努力を示されました。",
        15,
        after=8,
    )

    p = cell.add_paragraph()
    add_centered(
        p,
        "その成果をたたえ、ここに賞します。",
        18,
        bold=True,
        after=22,
    )

    if seal:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(seal), width=Cm(2.25))
        set_paragraph_spacing(p, before=0, after=8)

    p = cell.add_paragraph()
    add_centered(p, "令和　　　年　　　月　　　日", 13, font=FONT_SANS, after=4)

    p = cell.add_paragraph()
    add_centered(p, "ケニーズ学童クラブ", 15, color=DARK, bold=True, after=0)

    p = cell.add_paragraph()
    add_centered(p, "認定者　　　　　　　　　　　　　　　　印", 12, font=FONT_SANS, after=0)

    doc.save(DOCX_PATH)


def register_pdf_fonts():
    candidates = [
        ("YuMincho", "C:/Windows/Fonts/yumin.ttf"),
        ("YuMincho", "C:/Windows/Fonts/yumincho.ttc"),
        ("YuGothic", "C:/Windows/Fonts/yugothm.ttc"),
        ("YuGothicBold", "C:/Windows/Fonts/yugothb.ttc"),
    ]
    registered = set()
    for name, path in candidates:
        if name in registered or not Path(path).exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont(name, path))
            registered.add(name)
        except Exception:
            pass
    return {
        "serif": "YuMincho" if "YuMincho" in registered else "Helvetica",
        "sans": "YuGothic" if "YuGothic" in registered else "Helvetica",
        "sans_bold": "YuGothicBold" if "YuGothicBold" in registered else "Helvetica-Bold",
    }


def draw_centered_text(c, text, y, font, size, color_hex=DARK):
    c.setFillColor(colors.HexColor(f"#{color_hex}"))
    c.setFont(font, size)
    c.drawCentredString(landscape(A4)[0] / 2, y, text)


def build_pdf():
    PDF_OUT_DIR.mkdir(parents=True, exist_ok=True)
    seal = create_seal()
    fonts = register_pdf_fonts()
    width, height = landscape(A4)
    c = canvas.Canvas(str(PDF_PATH), pagesize=landscape(A4))

    c.setFillColor(colors.HexColor("#FFFDF4"))
    c.rect(0, 0, width, height, fill=1, stroke=0)

    margin = 14 * mm
    c.setStrokeColor(colors.HexColor(f"#{GOLD}"))
    c.setLineWidth(3.2)
    c.rect(margin, margin, width - margin * 2, height - margin * 2, fill=0, stroke=1)
    c.setLineWidth(0.9)
    c.rect(margin + 5 * mm, margin + 5 * mm, width - (margin + 5 * mm) * 2, height - (margin + 5 * mm) * 2, fill=0, stroke=1)

    draw_centered_text(c, "賞　状", height - 52 * mm, fonts["serif"], 43, DARK)
    draw_centered_text(c, "タイピング検定　中級", height - 72 * mm, fonts["sans_bold"], 20, GOLD)
    draw_centered_text(c, "受賞者名", height - 94 * mm, fonts["sans"], 13, "6E6255")
    draw_centered_text(c, "　　　　　　　　　　　　　　殿", height - 112 * mm, fonts["serif"], 23, DARK)
    c.setStrokeColor(colors.HexColor(f"#{GOLD}"))
    c.setLineWidth(1.1)
    c.line(width / 2 - 58 * mm, height - 116 * mm, width / 2 + 58 * mm, height - 116 * mm)

    draw_centered_text(
        c,
        "あなたはタイピング検定 中級において、正確で安定した入力技術と継続した努力を示されました。",
        height - 137 * mm,
        fonts["serif"],
        15,
        DARK,
    )
    draw_centered_text(c, "その成果をたたえ、ここに賞します。", height - 153 * mm, fonts["serif"], 18, DARK)

    if seal and seal.exists():
        c.drawImage(str(seal), width / 2 + 45 * mm, 24 * mm, 25 * mm, 25 * mm, mask="auto")

    draw_centered_text(c, "令和　　　年　　　月　　　日", 45 * mm, fonts["sans"], 13, DARK)
    draw_centered_text(c, "ケニーズ学童クラブ", 34 * mm, fonts["sans_bold"], 15, DARK)
    draw_centered_text(c, "認定者　　　　　　　　　　　　　　　　印", 24 * mm, fonts["sans"], 12, DARK)

    c.save()
    print(PDF_PATH)


def patch_page_background_and_geometry(source, target):
    with zipfile.ZipFile(source, "r") as zin:
        names = zin.namelist()
        content = {name: zin.read(name) for name in names}
    document_xml = content["word/document.xml"]
    root = ET.fromstring(document_xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    sect = root.find(".//w:sectPr", ns)
    pg_sz = sect.find("w:pgSz", ns)
    pg_sz.set(qn("w:w"), "16838")
    pg_sz.set(qn("w:h"), "11906")
    pg_sz.set(qn("w:orient"), "landscape")
    pg_mar = sect.find("w:pgMar", ns)
    pg_mar.set(qn("w:top"), "765")
    pg_mar.set(qn("w:right"), "850")
    pg_mar.set(qn("w:bottom"), "765")
    pg_mar.set(qn("w:left"), "850")
    background = ET.Element(f"{{{ns['w']}}}background")
    background.set(f"{{{ns['w']}}}color", "FFFDF4")
    root.insert(0, background)
    content["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in names:
            zout.writestr(name, content[name])


if __name__ == "__main__":
    build_docx()
    patch_page_background_and_geometry(DOCX_PATH, FINAL_PATH)
    build_pdf()
    print(FINAL_PATH)
